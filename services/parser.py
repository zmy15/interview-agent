"""文档解析服务 — PDF / Word / 代码文件 / 压缩包"""

import io
import logging
import os
import zipfile
import tarfile
import tempfile
import shutil
from pathlib import Path

logger = logging.getLogger(__name__)

# 延迟导入，避免缺失依赖时阻止模块加载


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """使用 PyMuPDF (fitz) 解析 PDF"""
    import fitz  # pymupdf
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    for page in doc:
        text = page.get_text()
        if text:
            text_parts.append(text)
    doc.close()
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    解析 DOCX 文档。
    先用 python-docx 提取结构化文本（段落/表格/页眉页脚），
    再用 XML 正则做全量提取，合并去重后返回最完整的结果。
    """
    from docx import Document

    # === 第1步：python-docx 结构化提取 ===
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # 段落文本
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # 表格文本
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                if cell.text.strip():
                    row_texts.append(cell.text.strip())
            if row_texts:
                parts.append(" | ".join(row_texts))

    # 页眉页脚（含表格）
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is None:
                continue
            for para in hf.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())

    structured = "\n".join(parts).strip()

    # === 第2步：XML 正则全量提取（捕获 python-docx 遗漏的文本） ===
    raw_text = _extract_docx_xml_fallback(file_bytes)

    # === 第3步：合并结果 ===
    if not structured:
        return raw_text
    if not raw_text:
        return structured

    # 如果 XML 提取的文本明显多于结构化文本，说明 python-docx 遗漏了内容
    # 用 XML 结果（纯文本拼接，丢失段落结构但保证完整）
    if len(raw_text) > len(structured) * 1.2:
        logger.info(
            f"XML fallback extracted more text ({len(raw_text)} chars) "
            f"than python-docx ({len(structured)} chars), using XML result"
        )
        return raw_text

    return structured


def _extract_docx_xml_fallback(file_bytes: bytes) -> str:
    """
    直接解析 DOCX 内部 XML，用正则提取所有 <w:t> 文本节点。
    不依赖 XML 树结构，能处理 AlternateContent、文本框、深层嵌套等复杂文档。
    这是 python-docx 返回空文本时的最后回退方案。
    """
    import re
    import xml.etree.ElementTree as ET

    # 优先查找的 XML 文件列表（按优先级）
    xml_parts = [
        "word/document.xml",
        "word/header1.xml", "word/header2.xml", "word/header3.xml",
        "word/footer1.xml", "word/footer2.xml", "word/footer3.xml",
        "word/footnotes.xml", "word/endnotes.xml",
    ]

    all_texts = []

    with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
        # 方法1：正则直接匹配 <w:t> 文本（最鲁棒，不依赖 XML 结构）
        for xml_part in xml_parts:
            if xml_part not in zf.namelist():
                continue
            try:
                xml_content = zf.read(xml_part).decode("utf-8", errors="replace")
                # 匹配 <w:t ...>text</w:t> 或 <w:t xml:space="preserve">text</w:t>
                texts = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', xml_content)
                all_texts.extend(t for t in texts if t)
            except Exception as e:
                logger.warning(f"Regex extraction failed for {xml_part}: {e}")

        # 方法2：如果正则也没提取到，用 ElementTree 遍历
        if not all_texts:
            for xml_part in xml_parts:
                if xml_part not in zf.namelist():
                    continue
                try:
                    xml_content = zf.read(xml_part)
                    root = ET.fromstring(xml_content)
                    for elem in root.iter():
                        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                        if tag == "t" and elem.text:
                            all_texts.append(elem.text)
                except Exception:
                    continue

        # 方法3：最后的尝试 — 遍历所有 XML 文件
        if not all_texts:
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                try:
                    xml_content = zf.read(name).decode("utf-8", errors="replace")
                    texts = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', xml_content)
                    all_texts.extend(t for t in texts if t)
                except Exception:
                    continue

    return "".join(all_texts)


def extract_text_from_doc(file_bytes: bytes) -> str:
    """
    解析旧版 .doc 文件（OLE 格式）。
    依赖 antiword 或 catdoc 系统工具，若无则尝试用 python-docx 解析（可能失败）。
    """
    # 尝试用 antiword
    import subprocess
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        result = subprocess.run(
            ["antiword", tmp_path], capture_output=True, text=True, timeout=30
        )
        os.unlink(tmp_path)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except (FileNotFoundError, Exception):
        pass

    # 回退：尝试 catdoc
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        result = subprocess.run(
            ["catdoc", tmp_path], capture_output=True, text=True, timeout=30
        )
        os.unlink(tmp_path)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except (FileNotFoundError, Exception):
        pass

    # 最后的回退：尝试当作文本读取（OLE 格式中可能嵌有可读文本）
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception:
        raise ValueError(
            "无法解析 .doc 文件（旧版 Word 格式）。"
            "请将文件另存为 .docx 格式后重新上传，"
            "或在系统中安装 antiword/catdoc 工具。"
        )


def extract_text_from_txt(file_bytes: bytes) -> str:
    """解析纯文本文件，自动处理编码"""
    # 优先尝试常见中文编码，UTF-16 放在最后以避免误识别
    for encoding in ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]:
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


def extract_code(file_bytes: bytes) -> str:
    """解析代码文件"""
    return extract_text_from_txt(file_bytes)


# ============ 项目/压缩包解析 ============

# 项目文件分类规则
PROJECT_FILE_CATEGORIES = {
    "source": {
        ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".cpp", ".c",
        ".h", ".hpp", ".rs", ".rb", ".php", ".swift", ".kt", ".scala",
        ".cs", ".vue", ".html", ".css", ".scss", ".less", ".sass",
        ".sql", ".sh", ".bash", ".ps1", ".bat", ".cmd", ".r", ".m",
    },
    "config": {
        ".yaml", ".yml", ".json", ".xml", ".toml", ".ini", ".cfg", ".conf",
        ".env", ".properties", ".editorconfig", ".prettierrc", ".eslintrc",
    },
    "document": {
        ".md", ".txt", ".rst", ".adoc", ".tex", ".pdf", ".docx", ".doc",
    },
    "build": {
        "Dockerfile", "Makefile", "CMakeLists.txt", ".dockerignore",
        ".gitignore", "docker-compose.yml", "docker-compose.yaml",
        "package.json", "Cargo.toml", "go.mod", "pom.xml", "build.gradle",
        "requirements.txt", "setup.py", "setup.cfg", "pyproject.toml",
    },
}

# 扩展名 → 分类
_EXT_TO_CATEGORY = {}
for _cat, _exts in PROJECT_FILE_CATEGORIES.items():
    for _ext in _exts:
        _EXT_TO_CATEGORY[_ext] = _cat


def _classify_file(filename: str) -> str:
    """根据文件名判断项目文件分类"""
    basename = os.path.basename(filename)
    ext = Path(filename).suffix.lower()

    # 特殊文件名匹配（无扩展名或特殊名称）
    if basename in PROJECT_FILE_CATEGORIES["build"]:
        return "build"

    # 以 test_ 开头或 _test 结尾的文件视为测试
    if basename.startswith("test_") or "_test" in basename:
        return "test"

    if ext in _EXT_TO_CATEGORY:
        return _EXT_TO_CATEGORY[ext]

    return "other"


def _is_text_file(filename: str) -> bool:
    """判断文件是否可能为文本文件（用于代码/文档提取）"""
    ext = Path(filename).suffix.lower()
    text_exts = (
        PROJECT_FILE_CATEGORIES["source"]
        | PROJECT_FILE_CATEGORIES["config"]
        | PROJECT_FILE_CATEGORIES["document"]
        | {".log", ".csv", ".tsv"}
    )
    if ext in text_exts:
        return True
    basename = os.path.basename(filename)
    return basename in PROJECT_FILE_CATEGORIES["build"]


def extract_archive(file_bytes: bytes, filename: str) -> dict[str, bytes]:
    """
    解压压缩包，返回 {内部路径: 文件字节} 字典。
    支持 zip / tar.gz / tar.bz2 / tar / 7z（如有 py7zr）。
    """
    ext = Path(filename).suffix.lower()
    namelower = filename.lower()
    extracted: dict[str, bytes] = {}

    if ext == ".zip" or namelower.endswith(".zip"):
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            for member in zf.infolist():
                if member.is_dir():
                    continue
                # 跳过 macOS 元数据（保留 .gitignore 等重要 dotfiles）
                basename = os.path.basename(member.filename)
                if "__MACOSX" in member.filename:
                    continue
                if basename.startswith("._") or basename == ".DS_Store":
                    continue
                try:
                    extracted[member.filename] = zf.read(member)
                except Exception as e:
                    logger.warning(f"跳过无法读取的文件 {member.filename}: {e}")

    elif namelower.endswith((".tar.gz", ".tgz", ".tar.bz2", ".tar.xz", ".tar")):
        mode_map = {
            ".tar.gz": "r:gz", ".tgz": "r:gz",
            ".tar.bz2": "r:bz2", ".tar.xz": "r:xz", ".tar": "r:",
        }
        mode = "r:*"
        for suffix, m in mode_map.items():
            if namelower.endswith(suffix):
                mode = m
                break
        with tarfile.open(fileobj=io.BytesIO(file_bytes), mode=mode) as tf:
            for member in tf.getmembers():
                if not member.isfile():
                    continue
                basename = os.path.basename(member.name)
                if "__MACOSX" in member.name:
                    continue
                if basename.startswith("._") or basename == ".DS_Store":
                    continue
                try:
                    f = tf.extractfile(member)
                    if f:
                        extracted[member.name] = f.read()
                except Exception as e:
                    logger.warning(f"跳过无法读取的文件 {member.name}: {e}")

    elif ext in (".7z",):
        try:
            import py7zr
            with py7zr.SevenZipFile(io.BytesIO(file_bytes)) as szf:
                for name, bio in szf.readall().items():
                    if name.endswith("/"):
                        continue
                    extracted[name] = bio.getvalue() if hasattr(bio, "getvalue") else bio.read()
        except ImportError:
            raise ValueError("解析 .7z 文件需要安装 py7zr 库，请执行: pip install py7zr")

    else:
        raise ValueError(f"不支持的压缩格式: {ext}，支持 .zip / .tar.gz / .tar.bz2 / .tar")

    return extracted


def analyze_project(file_bytes: bytes, filename: str) -> dict:
    """
    分析项目压缩包：解压 → 分类文件 → 提取文本 → 返回结构化结果。

    返回:
    {
        "filename": "project.zip",
        "structure": {"source": [...], "config": [...], "document": [...], "build": [...], "test": [...], "other": [...]},
        "file_count": 15,
        "total_text": "所有文本文件合并后的内容",
        "file_texts": {"src/main.py": "...", ...},
        "tech_stack": ["Python", "FastAPI", ...],  # 推断的技术栈
    }
    """
    extracted = extract_archive(file_bytes, filename)
    if not extracted:
        raise ValueError("压缩包为空或无法解压")

    structure: dict[str, list[str]] = {
        "source": [], "config": [], "document": [],
        "build": [], "test": [], "other": [],
    }
    file_texts: dict[str, str] = {}
    all_text_parts: list[str] = []
    tech_stack: set[str] = set()

    for internal_path, file_content in extracted.items():
        category = _classify_file(internal_path)
        structure[category].append(internal_path)

        if _is_text_file(internal_path):
            try:
                text = extract_text_from_txt(file_content)
                if text.strip():
                    file_texts[internal_path] = text
                    # 添加文件分隔标记，方便后续分块
                    all_text_parts.append(f"--- {internal_path} ---\n{text}")
            except Exception as e:
                logger.warning(f"无法解析文件 {internal_path}: {e}")

    # 推断技术栈
    tech_stack = _detect_tech_stack(extracted, structure)

    return {
        "filename": filename,
        "structure": structure,
        "file_count": len(extracted),
        "file_texts": file_texts,
        "total_text": "\n\n".join(all_text_parts),
        "tech_stack": sorted(tech_stack),
    }


def _detect_tech_stack(extracted: dict[str, bytes], structure: dict[str, list[str]]) -> set[str]:
    """根据项目文件推断技术栈"""
    stack: set[str] = set()
    all_files = [f.lower() for f in extracted.keys()]
    all_files_str = " ".join(all_files)

    # Python
    if any(f.endswith(".py") for f in all_files) or "requirements.txt" in all_files or "setup.py" in all_files or "pyproject.toml" in all_files:
        stack.add("Python")
        if "fastapi" in all_files_str or "fastapi" in str(extracted.get("requirements.txt", b"")):
            stack.add("FastAPI")
        if "flask" in all_files_str or "flask" in str(extracted.get("requirements.txt", b"")):
            stack.add("Flask")
        if "django" in all_files_str:
            stack.add("Django")

    # JavaScript / TypeScript
    if "package.json" in extracted:
        try:
            import json
            pkg = json.loads(extracted["package.json"])
            deps = {**pkg.get("dependencies", {}), **pkg.get("devDependencies", {})}
            if "react" in deps:
                stack.add("React")
            if "vue" in deps:
                stack.add("Vue")
            if "next" in deps:
                stack.add("Next.js")
            if "express" in deps:
                stack.add("Express")
            if "typescript" in deps or any(f.endswith(".ts") or f.endswith(".tsx") for f in all_files):
                stack.add("TypeScript")
            if any(f.endswith(".js") or f.endswith(".jsx") for f in all_files):
                stack.add("JavaScript")
            if "vite" in deps:
                stack.add("Vite")
            if "tailwindcss" in deps:
                stack.add("TailwindCSS")
        except Exception:
            if any(f.endswith((".ts", ".tsx")) for f in all_files):
                stack.add("TypeScript")
            if any(f.endswith((".js", ".jsx")) for f in all_files):
                stack.add("JavaScript")

    # Java
    if any(f.endswith(".java") for f in all_files):
        stack.add("Java")
        if "pom.xml" in all_files:
            stack.add("Maven")
        if "build.gradle" in all_files:
            stack.add("Gradle")
        if "spring" in all_files_str:
            stack.add("Spring")

    # Go
    if any(f.endswith(".go") for f in all_files):
        stack.add("Go")
        if "go.mod" in all_files:
            stack.add("Go Modules")

    # Rust
    if any(f.endswith(".rs") for f in all_files):
        stack.add("Rust")
        if "Cargo.toml" in all_files:
            stack.add("Cargo")

    # C / C++
    if any(f.endswith((".c", ".cpp", ".cc", ".cxx", ".h", ".hpp")) for f in all_files):
        stack.add("C/C++")
        if "CMakeLists.txt" in all_files:
            stack.add("CMake")

    # Docker
    if "Dockerfile" in all_files or "docker-compose.yml" in all_files or "docker-compose.yaml" in all_files:
        stack.add("Docker")

    # Git
    if ".gitignore" in all_files:
        stack.add("Git")

    return stack


# ============ 文件扩展名定义 ============

RESUME_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".cpp", ".c",
    ".h", ".hpp", ".rs", ".rb", ".php", ".swift", ".kt", ".scala",
    ".cs", ".vue", ".html", ".css", ".scss", ".sql", ".sh", ".bash",
    ".yaml", ".yml", ".json", ".xml", ".md", ".r", ".m",
}
ARCHIVE_EXTENSIONS = {".zip", ".tar.gz", ".tgz", ".tar.bz2", ".tar.xz", ".tar", ".7z"}


def parse_file(file_bytes: bytes, filename: str) -> str:
    """
    统一文件解析入口。
    根据扩展名分发到不同的解析器。
    """
    ext = Path(filename).suffix.lower()
    namelower = filename.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    elif ext == ".doc":
        return extract_text_from_doc(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    elif ext in CODE_EXTENSIONS:
        return extract_code(file_bytes)
    else:
        # 尝试作为文本文件解析
        return extract_text_from_txt(file_bytes)
